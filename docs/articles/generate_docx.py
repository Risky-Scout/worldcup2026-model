"""
Generate Word (.docx) articles for wizardofodds.com.
Run: python3.10 docs/articles/generate_docx.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from pathlib import Path

OUT_DIR = Path(__file__).parent

# ── Colors ────────────────────────────────────────────────────────────────────
GOLD      = RGBColor(0xD4, 0xAF, 0x37)
DARK_NAVY = RGBColor(0x0A, 0x0E, 0x1A)
BLUE_ACC  = RGBColor(0x1A, 0x3A, 0x6B)
BODY_CLR  = RGBColor(0x11, 0x11, 0x11)
MUTED_CLR = RGBColor(0x55, 0x55, 0x55)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRY = RGBColor(0xF0, 0xF0, 0xF0)
MID_GRY   = RGBColor(0xDD, 0xDD, 0xDD)
FORMULA_B = RGBColor(0xEE, 0xF2, 0xFF)
CALLOUT_B = RGBColor(0xFD, 0xF8, 0xED)


def new_doc():
    doc = Document()
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    return doc


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(rgb)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            for k, v in kwargs[edge].items():
                tag.set(qn(f'w:{k}'), v)
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def para_keep_with_next(para):
    pPr = para._p.get_or_add_pPr()
    kwn = OxmlElement('w:keepWithNext')
    pPr.append(kwn)
    kl = OxmlElement('w:keepLines')
    pPr.append(kl)


def add_site_header(doc, article_title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("WIZARDOFODDS.COM")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = GOLD
    run2 = p.add_run(f"     |     {article_title}")
    run2.font.size = Pt(9)
    run2.font.color.rgb = MUTED_CLR
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'D4AF37')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    para_keep_with_next(p)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BODY_CLR
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED_CLR
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    para_keep_with_next(p)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BODY_CLR
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '3')
    bottom.set(qn('w:color'), 'D4AF37')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    para_keep_with_next(p)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE_ACC
    return p


def add_body(doc, text, space_after=Pt(7), justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = space_after
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_CLR
    return p


def add_body_rich(doc, segments, space_after=Pt(7), justify=True):
    """segments = list of (text, bold) tuples."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = space_after
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold in segments:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(10)
        run.font.color.rgb = BODY_CLR
    return p


def add_formula(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x20, 0x60)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EEF2FF')
    pPr.append(shd)
    return p


def add_callout(doc, label, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    lp = cell.add_paragraph()
    lp.paragraph_format.space_before = Pt(2)
    lp.paragraph_format.space_after  = Pt(2)
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(9)
    lr.font.color.rgb = GOLD
    bp = cell.add_paragraph()
    bp.paragraph_format.space_before = Pt(0)
    bp.paragraph_format.space_after  = Pt(4)
    bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    br = bp.add_run(text)
    br.font.size = Pt(9.5)
    br.font.color.rgb = RGBColor(0x33, 0x33, 0x55)
    br.italic = True
    for spare in cell.paragraphs[:1]:
        spare._element.getparent().remove(spare._element)
    set_cell_bg(cell, CALLOUT_B)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'right', 'bottom', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'none')
        tcBorders.append(tag)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:color'), 'D4AF37')
    tcBorders.append(left)
    tcPr.append(tcBorders)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ": ")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = BODY_CLR
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
        r2.font.color.rgb = BODY_CLR
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = BODY_CLR
    return p


def add_spacer(doc, pts=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)
    return p


def add_page_break(doc):
    doc.add_page_break()


def dark_table_header(tbl, headers):
    hdr_row = tbl.rows[0]
    for i, cell in enumerate(hdr_row.cells):
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(headers[i])
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = GOLD
        set_cell_bg(cell, DARK_NAVY)


def style_data_row(tbl, row_idx, alt=False):
    row = tbl.rows[row_idx]
    bg = LIGHT_GRY if alt else WHITE
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = BODY_CLR


def add_disclaimer(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    run = p.add_run(
        "All probabilities represent regulation time (90 minutes + stoppage time) only. "
        "Extra time and penalties are excluded. "
        "This article is for informational and educational purposes. Please gamble responsibly. "
        "WizardOfOdds.com"
    )
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED_CLR


# =============================================================================
# ARTICLE 1 -- How the Model Works
# =============================================================================

def build_article1():
    doc = new_doc()
    TITLE = "How the 2026 World Cup Prediction Model Works"

    add_site_header(doc, TITLE)
    add_spacer(doc, 6)
    add_h1(doc, TITLE)
    add_subtitle(doc, "A complete guide to the mathematics behind the joint score probability engine -- WizardOfOdds.com -- June 2026")

    # ── Introduction ──────────────────────────────────────────────────────────
    add_h2(doc, "Introduction")
    add_body(doc,
        "Predicting the outcome of a football match is an exercise in structured humility. "
        "A match produces roughly two or three goals, arriving from a chaotic sequence of "
        "decisions and individual brilliance that no model can fully capture. Even the best "
        "football models in the world -- the ones running inside sharp sportsbooks -- are "
        "wrong more often than they are right about any individual game.")
    add_body(doc,
        "What a well-built model can do is get the probabilities right over a large number "
        "of matches. Assign 30% probability to outcomes that happen 30% of the time. Be "
        "faster than the public market in incorporating new information. Identify specific "
        "markets where the bookmaker's price does not reflect the available evidence. That "
        "is a narrower ambition than predicting winners -- but it is the one that produces "
        "long-run results.")
    add_body(doc,
        "This model produces a full joint probability distribution over every possible final "
        "score for every 2026 FIFA World Cup match. It blends 12 independent signal sources "
        "into a composite team rating, runs six competing parametric models daily, anchors "
        "each match's goal totals against live bookmaker lines, reconciles the result against "
        "market probabilities, and applies a Hierarchical Bayesian blend before final "
        "calibration. The pipeline runs without human intervention 24 hours a day and FTP-"
        "uploads results to WizardOfOdds.com after every run.")

    # ── What a PMF Is ────────────────────────────────────────────────────────
    add_h2(doc, "What a Joint Score PMF Actually Is")
    add_body(doc,
        "Before describing the engine, it helps to understand precisely what it produces. "
        "A PMF -- Probability Mass Function -- is a complete assignment of probability to "
        "every possible discrete outcome. For a football match, the joint score PMF is a "
        "two-dimensional grid. Each cell (h, a) holds P(Home = h, Away = a): the probability "
        "that the home team scores exactly h goals and the away team scores exactly a goals "
        "in regulation time.")
    add_body(doc,
        "The grid runs from (0, 0) through 8 or 9 goals per team, plus an explicit tail-mass "
        "bucket for extreme scores beyond the boundary. Every cell including the tail sums to "
        "exactly 1.0. This is a hard constraint enforced at every stage of the pipeline.")
    add_body(doc,
        "Every betting market is simply a different question about this same grid. There is "
        "no separate model for Over/Under, no separate model for BTTS -- every number flows "
        "from one source of truth:")

    tbl = doc.add_table(rows=9, cols=2)
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Inches(1.8)
    tbl.columns[1].width = Inches(4.4)
    dark_table_header(tbl, ["Market", "How It Is Calculated from the PMF Grid"])
    rows_data = [
        ("Over 2.5 goals",      "Sum all cells where h + a >= 3"),
        ("Under 2.5 goals",     "Sum all cells where h + a <= 2"),
        ("Both Teams to Score", "All cells where h >= 1 AND a >= 1"),
        ("Home Win",            "All cells where h > a"),
        ("Draw",                "All cells on the main diagonal: (0,0), (1,1), (2,2), ..."),
        ("Away Win",            "All cells where a > h"),
        ("Correct score 2-1",   "The single cell P(h=2, a=1) -- one number read directly"),
        ("Home clean sheet",    "Sum entire column a=0: P(0,0) + P(1,0) + P(2,0) + ..."),
    ]
    for i, (mkt, calc) in enumerate(rows_data, 1):
        row = tbl.rows[i]
        row.cells[0].text = mkt
        row.cells[1].text = calc
        style_data_row(tbl, i, alt=(i % 2 == 0))
    add_spacer(doc, 6)
    add_callout(doc, "The Internal Consistency Guarantee",
        "There is no separate model for each market type. The PMF is the entire model, and "
        "all markets are different questions about the same distribution. This guarantees "
        "internal consistency -- a property that does not hold on sites that blend data from "
        "different sources or use different models for different markets.")

    # ── Step 1: Composite Prior ───────────────────────────────────────────────
    add_h2(doc, "Step 1: Rating Every Team -- The 12-Signal Composite Prior")
    add_body(doc,
        "The first task is assigning each of the 48 World Cup teams an attack lambda and "
        "a defense lambda -- the expected goals scored and conceded against an average "
        "opponent on a neutral pitch. No single rating system is good enough to trust "
        "entirely. Club data does not transfer cleanly to international football. FIFA "
        "rankings can lag months behind a team's actual form. Bookmaker odds contain genuine "
        "signal but can be distorted by public betting flow on marquee matches. The solution "
        "is a composite prior blended from twelve independent signal sources.")

    add_h3(doc, "The Twelve Signal Sources")
    add_body(doc, "All twelve sources are active for every prediction:", space_after=Pt(4))
    tbl2 = doc.add_table(rows=13, cols=4)
    tbl2.style = 'Table Grid'
    col_widths = [0.4, 1.6, 1.7, 0.65]
    for i, w in enumerate(col_widths):
        for cell in tbl2.columns[i].cells:
            cell.width = Inches(w)
    dark_table_header(tbl2, ["#", "Signal Source", "Primary Data", "Weight"])
    sources = [
        ("1",  "market_implied",       "BDL consensus 1X2 odds, SHIN no-vig",               "~20%"),
        ("2",  "futures_implied",       "Tournament outright futures odds",                   "~5%"),
        ("3",  "best_player_form",      "Top-player performance ratings per team",            "~5%"),
        ("4",  "fifa_ranking",          "March 2026 official FIFA ranking points",            "~10%"),
        ("5",  "qualifying",            "Attack/defense efficiency, Bayesian-shrunk",         "~8%"),
        ("6",  "penaltyblog_pi",        "Pi Ratings -- goal-margin dynamic rating",           "~15%"),
        ("7",  "penaltyblog_elo",       "Elo rating, home_field_advantage=100",               "~12%"),
        ("8",  "massey",               "Massey method ratings",                              "~5%"),
        ("9",  "confederation",         "Confederation strength adjustment",                  "~5%"),
        ("10", "tournament_wc2026",     "In-tournament Bayesian shrinkage on WC2026 results", "~5%"),
        ("11", "injuries",             "avg_rating x (OUT=1.0 / GTD=0.5) impact score",     "~5%"),
        ("12", "intl_poisson",          "Bivariate Poisson on 49,433 Kaggle intl matches",   "~15%"),
    ]
    for i, (num, src, data, wt) in enumerate(sources, 1):
        row = tbl2.rows[i]
        row.cells[0].text = num
        row.cells[1].text = src
        row.cells[2].text = data
        row.cells[3].text = wt
        style_data_row(tbl2, i, alt=(i % 2 == 0))
    add_spacer(doc, 8)

    add_h3(doc, "Signal 1: market_implied (BDL Consensus Odds, ~20% weight)")
    add_body(doc,
        "When bookmaker odds are available, they encode information no rating system can "
        "access: late team news, undisclosed injuries, sharp professional money, and the "
        "aggregate view of every serious analyst who has looked at the match. The model "
        "reverse-engineers what attack and defense lambdas would produce the bookmaker's "
        "observed 1X2 probabilities. SHIN normalization removes the bookmaker margin before "
        "any comparison. Up to 6 sportsbooks contribute: FanDuel, DraftKings, BetMGM, "
        "BetRivers, Caesars, and Fanatics. The market_weight parameter is fixed at 0.20 for "
        "CLV-independence mode -- preventing the model from over-fitting to the very prices "
        "it is trying to beat.")

    add_h3(doc, "Signal 2: futures_implied (Tournament Outright Futures, ~5%)")
    add_body(doc,
        "Tournament outright futures odds (winner, top-4, group winner) are converted to "
        "team strength priors using a Bradley-Terry decomposition. These odds embed long-run "
        "tournament expectations and provide useful signal for teams where match-level 1X2 "
        "odds are thin or unavailable -- for example, in early group-stage fixtures.")

    add_h3(doc, "Signal 3: best_player_form (Top-Player Performance, ~5%)")
    add_body(doc,
        "For each team, the performance rating of the top-rated player currently in the "
        "tournament squad is used as a proxy for the team's ceiling performance level. "
        "This captures the 'superstar factor' that pure team-level ratings miss: the "
        "presence of an elite in-form striker or creative midfielder shifts a team's "
        "attacking potential beyond what aggregate ratings show.")

    add_h3(doc, "Signal 4: fifa_ranking (~10% weight)")
    add_body(doc,
        "FIFA's official points system, converted to an attack lambda via a calibrated "
        "sigmoid mapping. Captures long-run international performance across all "
        "competitions. Weakness: updates infrequently. The model uses the March 2026 "
        "snapshot -- the last official pre-tournament update.")

    add_h3(doc, "Signal 5: qualifying (~8% weight)")
    add_body(doc,
        "Each team's attack and defense efficiency during their qualifying campaign, "
        "Bayesian-shrunk toward the confederation average using the n/(n+3) formula. "
        "A team that dominated its qualifying group carries a meaningfully different "
        "expected output than one that scraped through.")

    add_h3(doc, "Signals 6-7: penaltyblog_pi and penaltyblog_elo (~27% combined)")
    add_body(doc,
        "Pi Rating (goal-margin-sensitive, ~15%) and Elo (~12%), both computed using the "
        "penaltyblog library. Pi updates on actual goal margins -- a 4-0 win earns a larger "
        "boost than a 1-0 win -- making it highly responsive to genuine form shifts. Elo "
        "updates only on match result (win/draw/loss), providing stability that prevents "
        "the prior from swinging sharply on an unusual scoreline. penaltyblog_elo is "
        "parameterized with home_field_advantage=100, calibrated to international data.")

    add_h3(doc, "Signal 8: massey (~5%)")
    add_body(doc,
        "Massey ratings solve a least-squares system over the full history of international "
        "results to assign each team a strength rating. They are particularly useful for "
        "teams with limited recent match data, where Elo and Pi have small sample noise. "
        "The Massey signal acts as an additional stabilizer for data-sparse regions.")

    add_h3(doc, "Signal 9: confederation (~5%)")
    add_body(doc,
        "Historical World Cup attack averages by confederation: CONMEBOL 1.45, UEFA 1.35, "
        "CONCACAF 1.20, CAF 1.10, AFC 1.10, OFC 0.90. Acts as a soft floor preventing "
        "any team from receiving a lambda wildly inconsistent with their region's historical "
        "output. Provides meaningful signal only for data-sparse teams.")

    add_h3(doc, "Signal 10: tournament_wc2026 (In-Tournament Bayesian Shrinkage, ~5%)")
    add_body(doc,
        "Once WC2026 matches complete, results feed back into each team's rating via "
        "Bayesian shrinkage. For n completed WC2026 matches, the shrinkage weight is "
        "n/(n+3): one match contributes 25%, two matches 40%, three matches 50%. "
        "Adjustments are capped at +/-30%. When BallDontLie shot data is available, "
        "tournament ratios blend 40% actual goals with 60% expected goals (xG) to "
        "reduce small-sample noise. As of June 2026: 22 teams have active adjustments.")

    add_h3(doc, "Signal 11: injuries (Blueprint Injury Impact Score, ~5%)")
    add_body(doc,
        "For each team, an injury impact score is computed from the BallDontLie squad "
        "and injury report data:")
    add_formula(doc,
        "injury_impact = SUM over injured players of:\n"
        "    player_avg_rating x status_weight\n"
        "\n"
        "Where status_weight = 1.0 for OUT, 0.5 for GTD (game-time decision)\n"
        "\n"
        "Higher impact score -> larger downward adjustment to attack/defense lambdas")
    add_body(doc,
        "This captures the effect of missing star players that no pre-tournament rating "
        "system can account for. A team missing its first-choice striker has a measurably "
        "different expected attack rate than their composite rating implies.")

    add_h3(doc, "Signal 12: intl_poisson (International Bivariate Poisson, 15% blend)")
    add_body(doc,
        "A dedicated Bivariate Poisson model fitted from scratch on 49,433 international "
        "football matches sourced from the Kaggle international results dataset, spanning "
        "1872 through 2026. A 3-year exponential half-life decay weighting ensures "
        "recent results dominate while the full historical depth provides stable "
        "long-run team strength estimates for every nation regardless of how active "
        "their recent calendar has been.")
    add_body(doc,
        "The intl_poisson signal is particularly valuable for: (a) neutralizing home "
        "advantage effects -- the historical dataset allows precise estimation of neutral-"
        "venue attack and defense rates for all 48 teams; (b) providing independent "
        "validation of the live-market and Elo signals; and (c) anchoring predictions "
        "for low-data teams whose Elo and Pi ratings rest on thin sample bases. The "
        "signal is blended at 15% weight into the composite prior.")
    add_callout(doc, "Why 12 Sources Instead of Fewer?",
        "Each source captures information the others miss. Market odds encode late-breaking "
        "news. Elo encodes long-run outcome stability. Pi encodes recent goal-margin form. "
        "intl_poisson encodes deep historical attack/defense patterns. Massey stabilizes "
        "data-sparse nations. Futures encode tournament-level expectations. Injuries encode "
        "what no pre-tournament rating can know. The ensemble is more robust than any "
        "single source, and correlation between sources is low enough that each genuinely "
        "adds information.")

    add_h3(doc, "Host and Altitude Adjustments")
    add_body(doc,
        "USA, Canada, and Mexico each receive +0.10 to attack lambda and -0.10 to defense "
        "lambda as co-hosts. All other matches are treated as neutral venue. Three Mexican "
        "venues sit at elevations that materially reduce scoring for unacclimated players. "
        "Both teams' lambdas are scaled equally since neither side has an acclimatization "
        "advantage:")
    add_bullet(doc, "Estadio Azteca, Mexico City (2,230m): 0.93x multiplier -- approximately 7% scoring reduction")
    add_bullet(doc, "Estadio Akron, Guadalajara (1,560m): 0.97x multiplier -- approximately 3% reduction")
    add_bullet(doc, "Estadio BBVA, Monterrey (530m): no adjustment")

    add_h3(doc, "Dynamic WC_AVG Scaling")
    add_body(doc,
        "This World Cup is playing at a historically high scoring rate. The historical "
        "average from 2018 and 2022 was 1.30 goals per team per match. Through completed "
        "2026 matches, the observed rate is 1.455 goals per team per match. The model applies "
        "a scaling factor of 1.119x to all 52 team lambdas after the composite prior is "
        "built, ensuring predictions reflect the actual tournament environment rather than "
        "being anchored to historical baselines.")

    # ── Step 2: Per-Match Calibration Enhancements ───────────────────────────
    add_h2(doc, "Step 2: Per-Match Calibration Enhancements")
    add_body(doc,
        "After the composite team lambdas are established, each individual match receives "
        "a set of per-match adjustments derived from live bookmaker data, group standings, "
        "and shot quality metrics. These adjustments fine-tune the team-level prior to "
        "the specific context of each fixture.")

    add_h3(doc, "total_anchor: BDL Market Over/Under Line")
    add_body(doc,
        "For each match, the model pulls the Over/Under total line from BallDontLie "
        "bookmaker data. The median across all available vendors is taken (clipped to the "
        "range 1.5 -- 6.0) and used as a per-match expected total goals anchor. This "
        "anchor is blended with the parametric estimate, ensuring the model's expected "
        "total goals stays grounded in current market pricing. When no BDL O/U data is "
        "available, the model falls back to the tournament-wide average.")

    add_h3(doc, "home_team_total / away_team_total: Individual Team Totals")
    add_body(doc,
        "Where BDL markets carry individual team total lines (e.g., Home Over 1.5 / "
        "Away Over 0.5), the implied team goal expectations are extracted and blended "
        "at 50% weight into the composite lambda for each team. This provides match-"
        "specific market intelligence at the team level, beyond just the match total.")

    add_h3(doc, "pts_diff / gd_diff: Group Standings Adjustments")
    add_body(doc,
        "In the group stage, each team's current points and goal differential in the "
        "group standings are used to apply a small attack multiplier. A team top of "
        "their group with a +5 goal differential enters the match with a demonstrated "
        "recent form advantage that may not yet be fully reflected in the composite "
        "rating. The multiplier is applied as:")
    add_formula(doc,
        "attack_mult = 1.0 + (pts_diff / max_pts_diff) * 0.08\n"
        "             + (gd_diff / max_gd_diff) * 0.08\n"
        "\n"
        "Where diffs are computed as (team value - opponent value)")

    add_h3(doc, "WDL Form String")
    add_body(doc,
        "Each team's recent Win/Draw/Loss sequence is parsed and converted to a form "
        "score. This form score is blended 50/50 with the rating z-score to produce "
        "a final momentum-adjusted attack rating. A team on a W-W-W-W run receives "
        "a meaningful uplift even if their underlying rating has not yet updated.")

    add_h3(doc, "confederation_diff: Cross-Confederation Adjustment")
    add_body(doc,
        "When two teams from different confederations meet, a confederation strength "
        "differential is computed using historical World Cup attack averages. The "
        "stronger confederation's team receives a small attack multiplier:")
    add_formula(doc, "attack_mult = 1.0 +/- confederation_strength_diff * [0.02 to 0.05]")

    add_h3(doc, "venue_lambda_adj: Travel, Rest, and Stadium Capacity")
    add_body(doc,
        "Three venue-specific factors are combined into a single lambda adjustment "
        "(promoted from shadow mode to full production):")
    add_bullet(doc, "Haversine travel distance from each team's base camp to the match venue -- longer travel applies a small downward adjustment to the traveling team's attack lambda")
    add_bullet(doc, "Rest days since each team's last match -- fewer rest days reduces expected output")
    add_bullet(doc, "Stadium capacity as a proxy for home-crowd atmosphere effects")
    add_body(doc,
        "The three factors are combined multiplicatively and applied to each team's "
        "attack and defense lambdas independently.")

    add_h3(doc, "Bayesian Blend: HierarchicalBayesianGoalModel at 20%")
    add_body(doc,
        "After the parametric PMF is generated, a HierarchicalBayesianGoalModel "
        "(implemented via penaltyblog) is computed for each match and blended at 20% "
        "weight into the final PMF. This Bayesian model uses partial pooling across "
        "all teams in the tournament -- a team with few completed matches borrows "
        "strength from teams with similar rating profiles. The blend is:")
    add_formula(doc,
        "final_PMF = 0.80 * parametric_PMF + 0.20 * bayesian_PMF")

    add_h3(doc, "Adaptive Temperature: Dynamic Calibration Weighting")
    add_body(doc,
        "The calibration temperature T shifts adaptively based on the number of "
        "completed WC2026 matches:")
    add_formula(doc,
        "If completed_matches >= 24:\n"
        "    weight = 60% WC2026 data / 40% historical (2018+2022)\n"
        "\n"
        "If completed_matches < 24:\n"
        "    weight = 30% WC2026 data / 70% historical")
    add_body(doc,
        "This prevents the calibration from over-weighting a small 2026 sample early "
        "in the tournament while ensuring the model transitions to current-tournament "
        "calibration as data accumulates.")

    add_h3(doc, "xGOT: Shot Quality Correction")
    add_body(doc,
        "When BallDontLie shot quality data is available, each team's mean xGOT/xG "
        "ratio (expected goals on target divided by expected goals) is computed as a "
        "measure of shot quality. A team generating more dangerous shot locations "
        "receives a multiplicative upward correction to their attack lambda:")
    add_formula(doc,
        "attack_lambda *= (team_mean_xGOT / team_mean_xG) / league_average_ratio")

    add_h3(doc, "calib_rho: Dixon-Coles Correlation Parameter")
    add_body(doc,
        "The Dixon-Coles low-score correlation parameter rho is calibrated from "
        "WC2026 data and blended with a market-derived prior:")
    add_formula(doc,
        "calib_rho = 0.60 * prior_rho + 0.40 * market_implied_rho\n"
        "\n"
        "Current calibrated value: rho = -0.042\n"
        "(Negative rho slightly reduces the probability of 0-0 and 1-1 draws\n"
        " relative to independent Poisson)")

    # ── Step 3: Bivariate Poisson ─────────────────────────────────────────────
    add_h2(doc, "Step 3: How P(Home=h, Away=a) Is Calculated -- The Bivariate Poisson")
    add_body(doc,
        "This is the most technically interesting stage of the model -- the one that "
        "separates it from the standard approach found in most public football forecasters.")

    add_h3(doc, "The Standard Approach and Its Flaw")
    add_body(doc,
        "The naive approach assumes home and away goals are independent Poisson random "
        "variables. Under this assumption, the joint probability of a final score is simply "
        "the product of two independent probabilities:")
    add_formula(doc,
        "P(H=h, A=a) = P(H=h) * P(A=a)\n"
        "           = [e^(-lam_h) * lam_h^h / h!] * [e^(-lam_a) * lam_a^a / a!]")
    add_body(doc,
        "This is the textbook starting point and it is wrong in a subtle but important way. "
        "Goals are not independent. When a goal goes in at the 25th minute, everything "
        "changes -- the losing team pushes forward, the winning team may sit deeper, the "
        "tactical picture reshapes entirely. High-intensity, open matches tend to produce "
        "more goals for both teams simultaneously. The correlation between home and away "
        "goals in football is small but real and positive.")

    add_h3(doc, "The Bivariate Poisson: Three Latent Processes")
    add_body(doc,
        "The Bivariate Poisson model handles this properly. The mathematical intuition: "
        "instead of two independent processes, model the goals as arising from THREE "
        "independent Poisson processes:")
    add_bullet(doc, "Z1 ~ Poisson(lambda_1): goals generated by the home team independently")
    add_bullet(doc, "Z2 ~ Poisson(lambda_2): goals generated by the away team independently")
    add_bullet(doc, "Z3 ~ Poisson(lambda_3): a shared 'match intensity' component -- extra goals that tend to occur simultaneously in high-energy matches")
    add_body(doc, "Then: H = Z1 + Z3,   A = Z2 + Z3")
    add_body(doc,
        "Z3 is the key innovation. It creates positive correlation between home and away "
        "goals through a shared latent process. A match with high Z3 is one where both "
        "teams are pushed into an open, attacking game.")

    add_h3(doc, "The Joint Probability Formula")
    add_formula(doc,
        "P(H=h, A=a) = e^(-(lam_1 + lam_2 + lam_3))\n"
        "              * SUM_{k=0}^{min(h,a)}\n"
        "                  [lam_1^(h-k) / (h-k)!]\n"
        "                * [lam_2^(a-k) / (a-k)!]\n"
        "                * [lam_3^k     / k!    ]")
    add_body(doc,
        "When lambda_3 = 0, the formula reduces exactly to the independent Poisson product. "
        "The Bivariate Poisson IS the independent Poisson when the data does not support "
        "correlation -- it never breaks, it only adds information when it is there.")
    add_callout(doc, "Current Calibrated Value",
        "lambda_3 = 0.170, calibrated from completed WC2026 matches. "
        "Positive and meaningful -- this World Cup is showing exactly the mutual intensity "
        "correlation the model was built to capture. Cov(H, A) = lambda_3 = 0.170.")

    add_h3(doc, "The Six Parametric Competitors")
    add_body(doc,
        "The model runs six competing parametric models daily, selects the winner by "
        "negative log-likelihood on held-out WC2026 data, and uses the champion for all "
        "predictions. The winner is reselected daily as more matches accumulate:", space_after=Pt(4))

    tbl3 = doc.add_table(rows=7, cols=3)
    tbl3.style = 'Table Grid'
    for cell in tbl3.columns[0].cells: cell.width = Inches(0.4)
    for cell in tbl3.columns[1].cells: cell.width = Inches(1.8)
    for cell in tbl3.columns[2].cells: cell.width = Inches(4.0)
    dark_table_header(tbl3, ["#", "Model", "What It Captures"])
    models = [
        ("1", "Independent Poisson",   "Classical baseline. Goals independent. Useful benchmark."),
        ("2", "Dixon-Coles",           "Low-score correction via rho = -0.042. Currently calibrated from WC2026 data."),
        ("3", "Bivariate Poisson",     "Shared intensity lambda_3 = 0.170. Current log-loss champion."),
        ("4", "Weibull Copula",        "Heavier tails than Poisson. More high-scoring outliers."),
        ("5", "Negative Binomial",     "Overdispersion: variance > mean. Common in real football data."),
        ("6", "Zero-Inflated Poisson", "Excess 0-0 draws beyond what Poisson predicts."),
    ]
    for i, (num, name, desc) in enumerate(models, 1):
        row = tbl3.rows[i]
        row.cells[0].text = num
        row.cells[1].text = name
        row.cells[2].text = desc
        style_data_row(tbl3, i, alt=(i % 2 == 0))
    add_spacer(doc, 8)

    # ── Step 4: Market Reconciliation ────────────────────────────────────────
    add_h2(doc, "Step 4: Market Reconciliation (SLSQP Optimization)")
    add_body(doc,
        "Even the Bivariate Poisson cannot know about a goalkeeper injury announced 90 "
        "minutes before kickoff. Market reconciliation solves this. For each match, the "
        "model extracts fair market probabilities from all available bookmakers after SHIN "
        "normalization across as many market types as the API carries: 1X2, Over/Under 0.5 "
        "through 6.5, BTTS, Draw No Bet, Double Chance, and correct score lines.")
    add_body(doc,
        "A constrained optimization algorithm -- SLSQP (Sequential Least Squares "
        "Programming) -- then adjusts the parametric PMF grid to satisfy these market "
        "constraints while minimizing KL divergence from the original distribution. KL "
        "divergence measures how far the adjusted distribution has moved from the starting "
        "point, ensuring the optimizer only moves as far as the market evidence requires. "
        "Constraints are hard: all cell values must be non-negative and sum to 1.0.")
    add_body(doc,
        "After SLSQP convergence, the HierarchicalBayesianGoalModel is blended at 20% "
        "weight into the final PMF (see Step 2). When SLSQP converges, it dominates for "
        "virtually all matches -- currently only 1 non-convergence out of 63 matches.")

    # ── Step 5: Calibration ────────────────────────────────────────────────
    add_h2(doc, "Step 5: Calibration (Adaptive Temperature Scaling)")
    add_body(doc,
        "A model is calibrated when the probabilities it outputs match observed frequencies. "
        "When it says 30%, outcomes should occur 30% of the time. Calibration is distinct "
        "from accuracy -- an overconfident model assigns 80% to outcomes that only occur "
        "65% of the time.")
    add_body(doc,
        "The model uses adaptive temperature scaling: a single parameter T fitted by "
        "minimizing exact-score log loss on out-of-sample predictions. Weighting shifts "
        "from 30/70 (WC2026/historical) to 60/40 once 24 or more matches complete. "
        "The calibrated probability for each cell is:")
    add_formula(doc,
        "p_calibrated[h,a]  proportional to  p_raw[h,a] ^ (1/T)\n"
        "\n"
        "Then renormalized so all cells sum to 1.0.\n"
        "\n"
        "Current calibrated temperature: T = 1.089\n"
        "(T > 1 flattens the distribution, correcting for mild overconfidence)")
    add_body(doc,
        "Five calibration metrics are evaluated on out-of-sample predictions only: "
        "exact-score negative log-likelihood (primary), Ranked Probability Score "
        "for 1X2, multiclass Brier score, Expected Calibration Error, and the ignorance "
        "score.")

    # ── Step 6: Edge Screening ─────────────────────────────────────────────
    add_h2(doc, "Step 6: Edge Screening and Kelly Sizing")
    add_body(doc,
        "With a calibrated PMF for each match, the model compares its probability "
        "estimates against the bookmaker's no-vig prices. An edge exists when the model's "
        "estimate is higher than the market's:")
    add_formula(doc,
        "Edge = (Model probability - Market implied probability) / Market implied probability")
    add_body(doc,
        "A market is flagged as a value opportunity only when all three conditions hold "
        "simultaneously:")
    add_bullet(doc, "Edge >= 4% (minimum threshold for signal above estimation noise)")
    add_bullet(doc, "90% confidence interval lower bound still exceeds market implied probability. CI computed by varying each team's lambda by +/- 12% and measuring the resulting range of PMF probabilities.")
    add_bullet(doc, "Market implied probability > 2% (excludes thin liquidity markets)")
    add_body(doc, "Bets passing all three filters are sized using the Kelly criterion:")
    add_formula(doc,
        "Full Kelly:   f* = Edge / (Decimal odds - 1)\n"
        "Half Kelly:   f* / 2    [default -- recommended]\n"
        "Hard cap:     5% of bankroll regardless of computed value")

    # ── Step 7: CLV ───────────────────────────────────────────────────────
    add_h2(doc, "Step 7: Closing Line Value (CLV) -- The Edge Litmus Test")
    add_body(doc,
        "Counting wins and losses is a poor way to evaluate a prediction model. The "
        "industry-standard measure that strips luck away from skill is Closing Line "
        "Value, or CLV.")
    add_body(doc,
        "The closing line is the bookmaker's final odds immediately before kickoff -- "
        "specifically the no-vig probability after removing the bookmaker margin. This "
        "represents the market's best collective estimate of the true probability, having "
        "absorbed every piece of publicly available information.")
    add_body(doc,
        "A model that consistently beats the closing line across a large sample is not "
        "getting lucky. It is providing information the market was slow to price in. "
        "Consistently losing to the closing line means the market was consistently smarter "
        "than the model, and the positive edges were illusions.")
    add_body(doc,
        "For every scheduled match the model tracks 15 markets: the three 1X2 outcomes, "
        "BTTS Yes and No, Over/Under at 0.5, 1.5, 2.5, 3.5, 4.5, 5.5, and 6.5 goals. "
        "Closing odds are captured automatically at T-3 minutes before each kickoff via "
        "a dedicated pre-match pipeline. The Market X-Ray page surfaces CLV tracking "
        "alongside real-time edge analysis for all active matches.")

    # ── Live Model ────────────────────────────────────────────────────────
    add_h2(doc, "The Live In-Play Model")
    add_body(doc,
        "During a live match, the model runs a non-homogeneous hazard model where "
        "goal-scoring rate varies by match minute. Goal rates are below average in the "
        "opening minutes, rise through the mid-half, spike just after half-time, and "
        "peak in the final ten minutes. On top of this temporal baseline, several "
        "live enhancements are active:")

    add_h3(doc, "Score-State Multipliers")
    tbl4 = doc.add_table(rows=6, cols=3)
    tbl4.style = 'Table Grid'
    for cell in tbl4.columns[0].cells: cell.width = Inches(2.3)
    for cell in tbl4.columns[1].cells: cell.width = Inches(1.5)
    for cell in tbl4.columns[2].cells: cell.width = Inches(1.5)
    dark_table_header(tbl4, ["Score State", "Home Rate", "Away Rate"])
    mult_data = [
        ("Draw at minute 60+",          "x1.10",  "x1.10"),
        ("Home team losing by 1",        "x1.25",  "x1.05"),
        ("Home team losing by 2+",       "x1.40",  "x1.10"),
        ("Home team winning by 1",       "x0.90",  "x1.10"),
        ("Home team winning by 2+",      "x0.80",  "x1.15"),
    ]
    for i, (state, hm, aw) in enumerate(mult_data, 1):
        row = tbl4.rows[i]
        row.cells[0].text = state
        row.cells[1].text = hm
        row.cells[2].text = aw
        style_data_row(tbl4, i, alt=(i % 2 == 0))
    add_spacer(doc, 6)

    add_h3(doc, "xG Blend (Active When BDL xG Available)")
    add_body(doc,
        "When live expected goals data is available from BallDontLie, the model blends "
        "the live xG rate (60% weight) with the pre-game prior (40% weight) starting from "
        "minute 15. Before minute 15, live xG from a small number of shots is too noisy "
        "to be useful.")

    add_h3(doc, "Momentum Scaling: passes_final_third")
    add_body(doc,
        "Live stats from BallDontLie include passes_final_third -- the number of passes "
        "each team has completed into the opponent's final third of the pitch. This is "
        "wired directly into the live hazard model as a momentum signal:")
    add_formula(doc,
        "hazard_mult = 1.0 + (team_passes_final_third - opponent_passes_final_third)\n"
        "              / normalizer * 0.03\n"
        "\n"
        "Range: approximately +/-3% attack rate adjustment")

    add_h3(doc, "momentum_df: BDL Match Momentum API")
    add_body(doc,
        "The BallDontLie match momentum API provides a real-time momentum score for each "
        "team. This momentum signal is active in the live hazard model:")
    add_formula(doc,
        "Home team momentum advantage: +8% hazard scaling\n"
        "Away team momentum advantage: +5% hazard scaling\n"
        "(Asymmetric -- home advantage amplifies momentum effect)")

    # ── Pipeline ──────────────────────────────────────────────────────────
    add_h2(doc, "The Automated Pipeline: A Living System")
    add_body(doc,
        "Every stage described above runs automatically, without human intervention, "
        "around the clock. After every run, results are FTP-uploaded to WizardOfOdds.com:")
    add_bullet(doc, "Daily at 8:00 AM UTC (4:00 AM ET): Full retraining. All 12 signals refreshed, all six parametric models retrained, composite prior rebuilt, per-match adjustments applied, predictions regenerated for all upcoming fixtures, calibration metrics logged. Updated JSONs deployed.")
    add_bullet(doc, "Hourly: Pre-match odds refresh. Lighter refresh capturing odds movements, updating CLV records, re-running predictions if newly completed matches appeared since the daily run.")
    add_bullet(doc, "Every 2 minutes during match hours: Live snapshot. During active match windows the pipeline self-chains -- each completed run immediately triggers the next. Current score, minute, xG, momentum, and passes_final_third wired into the live hazard model.")
    add_bullet(doc, "T-3 minutes before kickoff: Closing odds capture. A dedicated watcher checks the schedule every 15 minutes, sleeps until exactly 3 minutes before kickoff, fetches final market odds across all 15 tracked markets, records closing probabilities for CLV tracking.")

    # ── Limitations ────────────────────────────────────────────────────────
    add_h2(doc, "Honest Limitations")
    add_bullet(doc,
        "Regulation time only. All probabilities cover 90 minutes plus stoppage time. "
        "Extra time and penalty shootouts are not modeled.")
    add_bullet(doc,
        "lambda_3 = 0.170 is an average across match types. A match in stoppage time "
        "with a team pressing for an equalizer has dramatically different dynamics than "
        "a settled 3-0. The pre-game model uses a global calibrated value; the live model "
        "accounts for this via score-state multipliers.")
    add_bullet(doc,
        "The calibration temperature T and dynamic WC_AVG scaling factor will stabilize "
        "as the tournament progresses. Any parameter estimated from a small sample carries "
        "meaningful uncertainty.")
    add_bullet(doc,
        "Lambda uncertainty is fixed at +/- 12% for all teams. Per-team confidence "
        "intervals would require a fully Bayesian treatment.")
    add_bullet(doc,
        "Odds move between prediction time and kickoff. Edge estimates are calculated at "
        "the time the pipeline runs. Always verify current odds at your book before acting.")
    add_bullet(doc,
        "Positive expected value does not guarantee positive returns in any finite sample. "
        "Even a well-calibrated model with genuine edge will experience losing streaks.")

    add_spacer(doc, 12)
    add_disclaimer(doc)
    return doc


# =============================================================================
# ARTICLE 2 -- Page Guide
# =============================================================================

def build_article2():
    doc = new_doc()
    TITLE = "A Guide to the WC 2026 Prediction Pages"

    add_site_header(doc, TITLE)
    add_spacer(doc, 6)
    add_h1(doc, TITLE)
    add_subtitle(doc, "What every number, chart, and indicator means -- WizardOfOdds.com -- June 2026")

    add_body(doc,
        "There are five pages in the WC 2026 prediction section. Each draws from the same "
        "underlying joint score PMF engine and 12-signal composite rating system. The PMF "
        "is a two-dimensional grid where each cell (h, a) holds P(Home = h, Away = a) -- "
        "the probability that the home team scores exactly h goals and the away team scores "
        "exactly a goals in regulation time. Every market on every page flows from this "
        "single grid. The numbers cannot contradict each other by construction.")

    # ── Page 1 ────────────────────────────────────────────────────────────────
    add_h2(doc, "Page 1 -- Pre-Match Predictions")
    add_body(doc,
        "URL: pre-match.html",
        space_after=Pt(8))
    add_body(doc,
        "This is the command center. Every World Cup match scheduled for today appears in "
        "the main table with the model's probability estimates, expected goals, and edge "
        "analysis against live bookmaker prices. The data refreshes automatically. A red "
        "banner appears if the underlying prediction file is more than four hours old.")

    add_h3(doc, "The KPI Cards")
    add_bullet(doc, "The number of regulation kickoffs scheduled for today's date in Eastern Time.", bold_prefix="Matches Today")
    add_bullet(doc, "Individual betting markets across all today's matches passing all three edge filters simultaneously: raw edge >= 4%, 90% CI lower bound above market no-vig price, and market implied > 2%. This is a strict filter -- the count is frequently zero or very small. That is correct behavior.", bold_prefix="Value Bets")
    add_bullet(doc, "The single largest edge percentage found across all today's markets, with the specific market and match identified. An edge of +12% means the model estimates that outcome is 12% more likely than the bookmaker's no-vig price implies.", bold_prefix="Best Edge")
    add_bullet(doc, "The average total expected goals across today's fixtures -- the model's best estimate of how many goals each match will produce under current conditions.", bold_prefix="Avg xG / Match")

    add_h3(doc, "The Bankroll Sizing Tool")
    add_body(doc,
        "Enter a bankroll amount and select a Kelly fraction. For every market passing all "
        "three edge filters, the tool computes a recommended dollar stake.")
    add_bullet(doc, "Theoretically optimal. Produces large drawdowns when edge estimates contain error. Not recommended unless confidence in the edge is very high.", bold_prefix="Full Kelly")
    add_bullet(doc, "Bet size divided by two. Substantially reduces variance while retaining most compounding advantage. The default and standard recommendation.", bold_prefix="Half Kelly")
    add_bullet(doc, "Conservative setting. Appropriate when acknowledging significant uncertainty in the model's estimates.", bold_prefix="Quarter Kelly")
    add_body(doc, "All three fractions are hard-capped at 5% of entered bankroll regardless of what the formula computes.")

    add_h3(doc, "The Match Table")
    add_bullet(doc, "Home vs away. All matches treated as neutral venue except USA, Canada, and Mexico which carry a co-host adjustment (+0.10 attack, -0.10 defense).", bold_prefix="Match")
    add_bullet(doc, "Three-segment bar showing Home Win (gold), Draw (gray), Away Win (blue) probabilities for regulation time. Derived by summing the appropriate cells of the joint PMF grid.", bold_prefix="1X2 Probability Bars")
    add_bullet(doc, "Probability that total regulation goals exceed 2.5 -- three or more goals scored. Sum of all PMF cells where home_goals + away_goals >= 3.", bold_prefix="O/U 2.5")
    add_bullet(doc, "Both Teams to Score. Sum of all cells where home_goals >= 1 AND away_goals >= 1.", bold_prefix="BTTS")
    add_bullet(doc, "The single most probable final scoreline and its probability -- the peak cell of the joint grid.", bold_prefix="Top Score")
    add_bullet(doc, "The model's expected goals for home and away separately after market reconciliation -- lambda_home and lambda_away.", bold_prefix="xG (H-A)")
    add_bullet(doc, "Highest-edge market for this match passing all three value filters. If no market passes, this cell is blank.", bold_prefix="Best Edge / Fair Odds")

    add_h3(doc, "The Expanded Row")
    add_body(doc, "Click any match row to reveal three additional panels:")
    add_bullet(doc, "All non-trivial scorelines ranked from most to least likely, with proportion bars. Raw joint PMF values read directly from the grid.", bold_prefix="Full Scoreline Distribution")
    add_bullet(doc, "Every market the engine has priced: 1X2, BTTS, Over/Under at every standard line from 0.5 through 6.5, Draw No Bet, Double Chance, Win to Nil, Asian Handicap, and team-level totals.", bold_prefix="All Markets")
    add_bullet(doc, "For each market: model probability, market no-vig implied probability, edge %, fair odds, and current market odds. Rows highlighted in gold have passed all three value filters.", bold_prefix="Edge Report")

    # ── Page 2 ────────────────────────────────────────────────────────────────
    add_h2(doc, "Page 2 -- PMF Distributions")
    add_body(doc,
        "URL: pmf-distributions.html",
        space_after=Pt(8))
    add_body(doc,
        "Where Page 1 compresses the model output into a single row per match, this page "
        "opens it up completely. Select a match using the navigation chips at the top and "
        "every chart updates immediately to show the full probability landscape.")

    add_h3(doc, "Chart 1 -- Joint Score PMF Heatmap")
    add_body(doc,
        "The heatmap is the model in its purest form. Every cell (h, a) shows the "
        "probability that the home team scores exactly h goals and the away team scores "
        "exactly a goals. Home goals increase along the vertical axis, away goals along "
        "the horizontal. The color scale uses the square root of each cell's probability "
        "relative to the grid maximum -- without this transformation, cells above 8% would "
        "dominate and everything else would appear uniformly dark. Every betting market is "
        "readable directly from this heatmap.")

    add_h3(doc, "Chart 2 -- Marginal Goal Distributions")
    add_body(doc,
        "Two bar charts showing the probability each team scores exactly k goals. The "
        "marginal for the home team is obtained by summing across all away goal counts:")
    add_formula(doc, "P(Home = h) = SUM over all a of  P(Home = h, Away = a)")
    add_body(doc,
        "A team with a very tall bar at k=0 is frequently expected to be shut out. The "
        "tallest bar for most group-stage matches sits at k=1 for both teams.")

    add_h3(doc, "Chart 3 -- Total Goals Distribution")
    add_body(doc,
        "Shows P(total goals = k) for k = 0, 1, 2, ... Computed by summing all joint PMF "
        "cells along each anti-diagonal where h + a = k. Standard over/under lines (0.5, "
        "1.5, 2.5, 3.5, 4.5, 5.5) appear as vertical dividers. The probability to the "
        "right of any divider is the Over probability for that line.")

    add_h3(doc, "Chart 4 -- Goal Difference Distribution")
    add_body(doc,
        "Centered on zero. Gold bars are home wins (positive difference), gray is draw, "
        "blue is away wins. The height of the gray bar relative to gold and blue tells you "
        "how likely a draw is given the specific attacking and defensive qualities of these "
        "two teams.")

    add_h3(doc, "Chart 5 -- Top 20 Most Likely Scorelines")
    add_body(doc,
        "A ranked bar chart of the 20 highest-probability individual final scores with "
        "exact percentages and fair American odds. The top cell in most group-stage matches "
        "carries between 12% and 20% probability. If your sportsbook offers better odds "
        "than the fair odds shown, you may have found a value opportunity.")

    add_h3(doc, "Model vs Market Comparison Panel")
    add_body(doc,
        "A side-by-side table comparing model probabilities to BDL market no-vig prices "
        "for all standard markets: 1X2, BTTS, O/U lines from 0.5 through 6.5. Each row "
        "shows the model probability, market-implied probability, edge, and fair odds. "
        "This is the same data surfaced in more detail on the Market X-Ray page.")

    add_h3(doc, "O/U Lines Table")
    add_body(doc,
        "Over and Under probabilities for every standard total line from 0.5 through 6.5, "
        "all in one place. Every number is computed by summing the appropriate region of "
        "the joint PMF -- no approximations.")

    # ── Page 3 ────────────────────────────────────────────────────────────────
    add_h2(doc, "Page 3 -- Live In-Play PMF")
    add_body(doc,
        "URL: live-pmf.html",
        space_after=Pt(8))
    add_body(doc,
        "This page activates when a World Cup match is in progress. When no match is live, "
        "it shows the next scheduled kickoff time. All probabilities are regulation time only. "
        "Data updates every ~5 minutes during live matches.")

    add_h3(doc, "How the Live Model Differs From Pre-Game")
    add_body(doc,
        "The pre-game model asks: what will the final score be from kickoff? The live model "
        "asks: given that the current score is H-A at minute t, what will the final score "
        "be at 90+ minutes? This is conditional probability. Any score lower than the "
        "current score is impossible -- those cells of the PMF grid are locked at zero. "
        "The remaining probability redistributes entirely across reachable scores.")
    add_body(doc,
        "The live model uses a non-homogeneous hazard model -- goal-scoring rate varies "
        "by match minute, calibrated from 2018 and 2022 World Cup data. On top of this "
        "temporal baseline, score-state multipliers adjust each team's rate, xG data from "
        "BallDontLie is blended when available (60% live / 40% prior from minute 15), "
        "passes_final_third wires momentum directly into the hazard (+-3% attack), and "
        "the BDL match momentum API applies +-8/5% hazard scaling when available.")

    add_h3(doc, "Score-State Multipliers")
    tbl5 = doc.add_table(rows=6, cols=2)
    tbl5.style = 'Table Grid'
    for cell in tbl5.columns[0].cells: cell.width = Inches(2.5)
    for cell in tbl5.columns[1].cells: cell.width = Inches(3.75)
    dark_table_header(tbl5, ["Score State", "Effect on Goal Rates"])
    mult_data2 = [
        ("Draw at minute 60+",         "Both teams x1.10 -- games open up in the final 30 minutes"),
        ("Home team losing by 1",       "Home x1.25, Away x1.05 -- counter-attack risk increases"),
        ("Home team losing by 2+",      "Home x1.40, Away x1.10"),
        ("Home team winning by 1",      "Home x0.90, Away x1.10 -- away team pushes forward"),
        ("Home team winning by 2+",     "Home x0.80, Away x1.15"),
    ]
    for i, (state, effect) in enumerate(mult_data2, 1):
        row = tbl5.rows[i]
        row.cells[0].text = state
        row.cells[1].text = effect
        style_data_row(tbl5, i, alt=(i % 2 == 0))
    add_spacer(doc, 6)

    add_h3(doc, "Connection Badge")
    add_bullet(doc, "Active push connection. When a goal or status change is reported, the server recomputes the full conditional PMF and pushes it to connected browsers. Target latency: under 200 milliseconds.", bold_prefix="WebSocket (green)")
    add_bullet(doc, "Push connection unavailable. The page fetches updated data from a static JSON file every 60 seconds. Automatic fallback -- no user action needed.", bold_prefix="Polling (yellow)")

    add_h3(doc, "Live KPI Cards")
    add_bullet(doc, "Number of World Cup matches currently in progress.", bold_prefix="Matches Live")
    add_bullet(doc, "Next scheduled match with Eastern Time kickoff.", bold_prefix="Next Kickoff")
    add_bullet(doc, "Total goals scored across all live and recently completed matches.", bold_prefix="Goals Today")
    add_bullet(doc, "Time since the last live snapshot. Under 2 minutes is normal. Above 10 minutes during a live match triggers a health warning banner.", bold_prefix="Data Age")

    add_h3(doc, "Win Probability Bar and Shift Table")
    add_body(doc,
        "The same three-segment bar as on Page 1 -- Home Win (gold), Draw (gray), Away Win "
        "(blue) -- but now conditional on the current score and minute. Directly below the "
        "bar, the Pre-Game to Live Shift table shows, for each main market, the pre-game "
        "probability, current live probability, and arithmetic difference.")

    add_h3(doc, "Win Probability Sparkline")
    add_body(doc,
        "A compact line chart showing the home team's win probability from kickoff to the "
        "current minute. Sharp upward jumps are home goals; sharp downward drops are away "
        "goals. This history accumulates within your browser session and resets on page reload.")

    add_h3(doc, "Live Joint Score PMF Heatmap")
    add_body(doc,
        "The same heatmap as on Page 2, updated with every live snapshot. The cell "
        "corresponding to the current live score is outlined in red. All unreachable cells "
        "are forced to zero and appear dark. As the match progresses, the dark region grows "
        "and probability concentrates into fewer and fewer bright cells. In stoppage time "
        "of a 1-0 match, the heatmap may assign 85-90% probability to the single 1-0 cell.")

    add_h3(doc, "Next Goal Probabilities")
    add_formula(doc,
        "Home scores next: lambda_h_rem / (lambda_h_rem + lambda_a_rem)\n"
        "Away scores next: lambda_a_rem / (lambda_h_rem + lambda_a_rem)\n"
        "No more goals:   e^(-lambda_h_rem) * e^(-lambda_a_rem)")
    add_body(doc,
        "As the match enters stoppage time with small remaining expected goals, the "
        "'No more goals' value typically climbs above 80-90%.")

    add_h3(doc, "Top 10 Most Likely Final Scores (Live)")
    add_body(doc,
        "Same ranked list as Page 2, restricted to reachable outcomes only. The current "
        "live score is marked. As the match approaches the final whistle, the probability "
        "on the leading score climbs rapidly.")

    # ── Page 4 ────────────────────────────────────────────────────────────────
    add_h2(doc, "Page 4 -- Live Pitch")
    add_body(doc,
        "URL: live-pitch.html",
        space_after=Pt(8))
    add_body(doc,
        "The Live Pitch page provides a real-time animated shot map for any match currently "
        "in progress, rendered using BallDontLie player coordinate data. It is the most "
        "visually intensive page in the suite and is designed for in-play monitoring.")

    add_h3(doc, "Animated Shot Map")
    add_body(doc,
        "Each shot attempt in the live match is plotted on a scaled pitch diagram using "
        "the player_x and player_y coordinates from the BallDontLie live player data feed. "
        "Shot markers are color-coded by outcome:")
    add_bullet(doc, "Gold filled circle: goal scored")
    add_bullet(doc, "White open circle: shot on target, saved")
    add_bullet(doc, "Gray X: shot off target or blocked")
    add_body(doc,
        "Markers animate onto the pitch at the minute they occurred, allowing you to "
        "reconstruct the attacking flow of the match visually. Older shots fade slightly "
        "to help distinguish recent activity from early match shots.")

    add_h3(doc, "Momentum KPIs")
    add_body(doc,
        "Below the pitch, a row of key performance indicators derived from the live stats "
        "feed shows each team's current in-match momentum. These same KPIs feed directly "
        "into the live hazard model:")
    add_bullet(doc, "Passes into the final third, updated each snapshot. The ratio between teams wires into the +/-3% attack hazard adjustment.", bold_prefix="Passes Final Third")
    add_bullet(doc, "The BDL match momentum score for each team, displayed as a bar. When one team's momentum score exceeds the threshold, the +8% (home) or +5% (away) hazard scaling activates.", bold_prefix="Match Momentum")
    add_bullet(doc, "Running xG for each team accumulated during the match, blended at 60% into the live hazard from minute 15 onward.", bold_prefix="Live xG")
    add_bullet(doc, "Total shots attempted by each team in the match.", bold_prefix="Shots")
    add_bullet(doc, "Shots on target for each team -- the most predictive single stat for in-play goal likelihood.", bold_prefix="Shots on Target")

    add_h3(doc, "Data Freshness")
    add_body(doc,
        "The Live Pitch updates on the same 2-minute snapshot cycle as the Live PMF page. "
        "A timestamp shows the age of the current coordinate data. Because BDL player "
        "coordinate data has slightly higher latency than score/clock data, the shot map "
        "may lag the Live PMF page by one snapshot cycle during fast-moving moments.")

    # ── Page 5 ────────────────────────────────────────────────────────────────
    add_h2(doc, "Page 5 -- Market X-Ray")
    add_body(doc,
        "URL: market-xray/index.html",
        space_after=Pt(8))
    add_body(doc,
        "The Market X-Ray is a trader-grade analysis tool providing the deepest level of "
        "model-vs-market comparison available on the site. It is designed for users who "
        "want to go beyond edge percentages and understand the full picture of value, "
        "confidence, and market movement for every active match.")

    add_h3(doc, "Fair Odds vs Market Comparison")
    add_body(doc,
        "For each match and each market, the X-Ray shows the model's fair (no-vig) odds "
        "alongside the current market odds from all available BDL bookmakers. The "
        "comparison is presented as both American odds and implied probability, making "
        "it easy to identify where the model and market diverge significantly.")

    add_h3(doc, "Edge, EV, and Confidence Grades")
    add_bullet(doc, "The percentage by which the model's probability exceeds the market's no-vig probability. Computed identically to the edge on Page 1 -- same formula, same source.", bold_prefix="Edge")
    add_bullet(doc, "Expected Value: the dollar return per $100 wagered assuming the model's probability is the true probability. EV = (model_prob x net_payout) - (1 - model_prob) x 100.", bold_prefix="EV")
    add_bullet(doc, "A letter grade (A through F) reflecting the combined strength of the edge, the CI lower bound check, and the number of bookmakers confirming the price. An A grade means the edge is large, robust to lambda uncertainty, and confirmed across multiple books.", bold_prefix="Confidence Grade")

    add_h3(doc, "Trader Action Notes")
    add_body(doc,
        "Each market receives one of six action labels derived from the edge and "
        "confidence grade:")
    tbl6 = doc.add_table(rows=7, cols=2)
    tbl6.style = 'Table Grid'
    for cell in tbl6.columns[0].cells: cell.width = Inches(1.4)
    for cell in tbl6.columns[1].cells: cell.width = Inches(4.8)
    dark_table_header(tbl6, ["Action", "Meaning"])
    actions = [
        ("BET",           "Edge >= 8%, Confidence A or B, CI lower bound clears market. Full signal -- act at current price."),
        ("SMALL BET",     "Edge 4-8%, Confidence B or C. Positive signal but uncertainty warrants reduced size."),
        ("LEAN",          "Edge 2-4%, not quite threshold. Monitor -- may develop into a BET as odds move."),
        ("WAIT",          "Edge present but line moving unfavorably. Do not act until line stabilizes."),
        ("PASS",          "Edge below threshold or CI check fails. No actionable value at current price."),
        ("DO NOT CHASE",  "Line has already moved significantly toward model fair value. CLV has been consumed."),
    ]
    for i, (act, meaning) in enumerate(actions, 1):
        row = tbl6.rows[i]
        row.cells[0].text = act
        row.cells[1].text = meaning
        style_data_row(tbl6, i, alt=(i % 2 == 0))
    add_spacer(doc, 6)

    add_h3(doc, "Line Movement Tracking")
    add_body(doc,
        "The X-Ray records the opening line, the current line, and the direction and "
        "magnitude of movement for each market. A sparkline shows the price history "
        "from the model's first prediction through the current moment. Markets moving "
        "toward the model's fair value confirm the signal; markets moving away suggest "
        "new information the model has not yet incorporated.")

    add_h3(doc, "CLV Tracker")
    add_body(doc,
        "For matches that have already kicked off, the CLV Tracker shows the model's "
        "opening prediction probability versus the closing no-vig probability for all "
        "15 tracked markets. Positive CLV (model was ahead of where the market closed) "
        "is highlighted in gold. Negative CLV is shown in muted gray. The aggregate "
        "CLV across all markets and all completed matches is displayed as the primary "
        "performance metric at the top of the tracker panel.")
    add_callout(doc, "How to Use the Market X-Ray",
        "Start with the Trader Action Notes. BET and SMALL BET are the only actionable "
        "signals. LEAN is worth monitoring. WAIT means check back after the next pipeline "
        "run. PASS and DO NOT CHASE mean no action at this price. The CLV Tracker "
        "accumulates over the tournament and is the fairest measure of whether the model's "
        "edges have been genuine.")

    # ── Limitations ────────────────────────────────────────────────────────
    add_h2(doc, "Scope and Limitations -- All Pages")
    add_bullet(doc, "All probabilities represent regulation time (90 minutes plus stoppage time) only. Extra time and penalty shootouts are not modeled.")
    add_bullet(doc, "The Bivariate Poisson substantially reduces the independence assumption (lambda_3 = 0.170), but it remains an average across match types. Score-state multipliers in the live model provide additional correction.")
    add_bullet(doc, "Calibration rests on 2018 and 2022 World Cup data augmented by completed 2026 matches. The adaptive temperature weighting (30/70 vs 60/40) shifts automatically as 2026 data accumulates.")
    add_bullet(doc, "Edge estimates are outputs of a probabilistic model. Market odds move between prediction time and kickoff. Always verify current prices at your book before acting.")
    add_bullet(doc, "WebSocket update speed is subject to network latency and server load. The stated target of under 200 milliseconds applies under normal conditions.")
    add_bullet(doc, "Live Pitch coordinate data (player_x/player_y) may lag the Live PMF score data by one snapshot cycle during fast-moving moments.")

    add_spacer(doc, 12)
    add_disclaimer(doc)
    return doc


# =============================================================================
# RENDER
# =============================================================================

if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    print("Generating Word documents...")

    doc1 = build_article1()
    p1 = OUT_DIR / "wc2026-how-the-model-works.docx"
    doc1.save(str(p1))
    print(f"  {p1.name}  ({p1.stat().st_size // 1024} KB)")

    doc2 = build_article2()
    p2 = OUT_DIR / "wc2026-page-guide.docx"
    doc2.save(str(p2))
    print(f"  {p2.name}  ({p2.stat().st_size // 1024} KB)")

    print("Done.")
